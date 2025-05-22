import os
import json
import csv
import pandas as pd
from flask import Flask, render_template, request, redirect, url_for, flash, send_file
from werkzeug.utils import secure_filename
from datetime import datetime
import uuid
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
import docx
from docx.shared import Inches
import numpy as np
import requests
from dotenv import load_dotenv

# Load environment variables from .env file
load_dotenv()

# Initialize Flask app
app = Flask(__name__)
app.secret_key = "ugc_budget_allocation_secret_key"

# Configure upload and report folders
UPLOAD_FOLDER = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'uploads')
REPORT_FOLDER = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'reports')
ALLOWED_EXTENSIONS = {'csv', 'json'}

# Groq API configuration
GROQ_API_KEY = os.environ.get("GROQ_API_KEY")
GROQ_API_URL = "https://api.groq.com/openai/v1/chat/completions"

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['REPORT_FOLDER'] = REPORT_FOLDER

# Ensure directories exist
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(REPORT_FOLDER, exist_ok=True)

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# Function to generate enhanced recommendations using Groq LLM
def generate_enhanced_recommendations(institution_data, basic_recommendations):
    """Generate detailed recommendations using Groq LLM API"""
    if not GROQ_API_KEY:
        return basic_recommendations  # Fallback to basic recommendations if no API key
    
    try:
        # Prepare the prompt for the LLM
        prompt = f"""
        As an educational budget allocation expert, provide detailed recommendations for improving the following institution:
        
        Institution Name: {institution_data.get('name', 'Unknown Institution')}
        
        Current Scores:
        - Infrastructure: {institution_data.get('infrastructure', 'N/A')}/10
        - Faculty: {institution_data.get('faculty', 'N/A')}/10
        - Research: {institution_data.get('research', 'N/A')}/10
        - Students: {institution_data.get('students', 'N/A')}/10
        - Placement: {institution_data.get('placement', 'N/A')}/10
        
        Basic Recommendations:
        {', '.join(basic_recommendations) if basic_recommendations else 'None'}
        
        Please provide 3-5 specific, actionable recommendations that would help this institution improve its scores and budget allocation. For each recommendation, include:
        1. A clear action item
        2. Expected impact on scores
        3. Implementation timeframe (short/medium/long term)
        
        Important: When mentioning any monetary values, always use "Rs." (Indian Rupees) instead of $ or other currency symbols.
        
        Format each recommendation as a bullet point.
        """
        
        # Call Groq API
        headers = {
            "Authorization": f"Bearer {GROQ_API_KEY}",
            "Content-Type": "application/json"
        }
        
        payload = {
            "model": "llama3-8b-8192",  # Using Llama 3 model
            "messages": [
                {"role": "system", "content": "You are an expert in higher education budget allocation and institutional improvement."},
                {"role": "user", "content": prompt}
            ],
            "temperature": 0.7,
            "max_tokens": 1024
        }
        
        response = requests.post(GROQ_API_URL, headers=headers, json=payload)
        response.raise_for_status()
        
        # Extract recommendations from response
        result = response.json()
        enhanced_recommendations = result["choices"][0]["message"]["content"].strip().split("\n")
        
        # Clean up the recommendations (remove empty lines, etc.)
        enhanced_recommendations = [rec.strip() for rec in enhanced_recommendations if rec.strip()]
        
        # Replace any $ symbols with Rs.
        enhanced_recommendations = [rec.replace('$', 'Rs.') for rec in enhanced_recommendations]
        
        return enhanced_recommendations
    
    except Exception as e:
        print(f"Error generating enhanced recommendations: {str(e)}")
        return basic_recommendations  # Fallback to basic recommendations

# UGC Budget Allocation Agent
class UGCBudgetAgent:
    def __init__(self):
        # Define weights for different criteria
        self.weights = {
            'infrastructure': 0.25,
            'faculty': 0.30,
            'research': 0.20,
            'students': 0.15,
            'placement': 0.10
        }
        
        # Define budget allocation rules
        self.total_budget = 10000000  # 10 million base budget
        self.min_allocation = 500000  # Minimum allocation per institution
    
    def analyze_data(self, data):
        """Analyze institution data and calculate budget allocations"""
        results = []
        
        # Calculate total score for each institution
        for institution in data:
            total_score = 0
            for criterion, weight in self.weights.items():
                if criterion in institution:
                    total_score += float(institution[criterion]) * weight
            
            # Calculate budget allocation based on score
            # Score is normalized to be between 0 and 1
            normalized_score = total_score / 10  # Assuming scores are out of 10
            
            # Calculate allocation (min allocation + performance-based allocation)
            allocation = self.min_allocation + (normalized_score * (self.total_budget - self.min_allocation * len(data)))
            
            # Generate basic recommendations
            basic_recommendations = self.generate_recommendations(institution)
            
            # Generate enhanced recommendations using LLM
            enhanced_recommendations = generate_enhanced_recommendations(institution, basic_recommendations)
            
            results.append({
                'name': institution.get('name', 'Unknown Institution'),
                'total_score': round(total_score, 2),
                'budget_allocation': round(allocation, 2),
                'recommendations': enhanced_recommendations
            })
        
        return results
    
    def generate_recommendations(self, institution):
        """Generate basic recommendations based on institution scores"""
        recommendations = []
        
        # Check each criterion and provide recommendations for improvement
        for criterion, weight in self.weights.items():
            if criterion in institution:
                score = float(institution[criterion])
                if score < 5:
                    recommendations.append(f"Improve {criterion} (current score: {score}/10)")
                elif score < 7:
                    recommendations.append(f"Consider enhancing {criterion} (current score: {score}/10)")
        
        return recommendations

# Initialize the UGC Budget Agent
ugc_agent = UGCBudgetAgent()

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        flash('No file part')
        return redirect(request.url)
    
    file = request.files['file']
    
    if file.filename == '':
        flash('No selected file')
        return redirect(request.url)
    
    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)
        
        # Process the file
        try:
            data = process_file(file_path)
            results = ugc_agent.analyze_data(data)
            
            # Generate report
            report_format = request.form.get('report_format', 'pdf')
            report_path = generate_report(results, report_format)
            
            return redirect(url_for('download_report', filename=os.path.basename(report_path)))
        
        except Exception as e:
            flash(f'Error processing file: {str(e)}')
            return redirect(url_for('index'))
    
    flash('Invalid file type. Please upload CSV or JSON files only.')
    return redirect(url_for('index'))

def process_file(file_path):
    """Process uploaded file and extract institution data"""
    data = []
    
    if file_path.endswith('.csv'):
        with open(file_path, 'r') as csvfile:
            reader = csv.DictReader(csvfile)
            for row in reader:
                data.append(row)
    
    elif file_path.endswith('.json'):
        with open(file_path, 'r') as jsonfile:
            data = json.load(jsonfile)
    
    return data

def generate_report(results, format='pdf'):
    """Generate budget allocation report in the specified format"""
    timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
    report_id = f"{timestamp}_{uuid.uuid4().hex[:8]}"
    
    if format.lower() == 'pdf':
        return generate_pdf_report(results, report_id)
    else:
        return generate_doc_report(results, report_id)

def generate_pdf_report(results, report_id):
    """Generate PDF report with budget allocation results"""
    report_path = os.path.join(app.config['REPORT_FOLDER'], f"budget_report_{report_id}.pdf")
    
    doc = SimpleDocTemplate(report_path, pagesize=letter)
    styles = getSampleStyleSheet()
    elements = []
    
    # Add title
    title = Paragraph("UGC Budget Allocation Report", styles['Title'])
    elements.append(title)
    elements.append(Spacer(1, 12))
    
    # Add date
    date_text = Paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", styles['Normal'])
    elements.append(date_text)
    elements.append(Spacer(1, 24))
    
    # Add summary
    summary = Paragraph(f"Total institutions analyzed: {len(results)}", styles['Heading2'])
    elements.append(summary)
    elements.append(Spacer(1, 12))
    
    # Create table for budget allocations
    data = [["Institution", "Score", "Budget Allocation (Rs.)"]]
    for result in results:
        data.append([
            result['name'],
            f"{result['total_score']}/10",
            f"Rs. {result['budget_allocation']:,.2f}"
        ])
    
    table = Table(data, colWidths=[250, 100, 150])
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)
    ]))
    
    elements.append(table)
    elements.append(Spacer(1, 24))
    
    # Add detailed recommendations for each institution
    elements.append(Paragraph("Detailed Recommendations", styles['Heading2']))
    elements.append(Spacer(1, 12))
    
    for result in results:
        elements.append(Paragraph(f"Institution: {result['name']}", styles['Heading3']))
        elements.append(Spacer(1, 6))
        
        if result['recommendations']:
            for recommendation in result['recommendations']:
                elements.append(Paragraph(f"• {recommendation}", styles['Normal']))
        else:
            elements.append(Paragraph("No specific recommendations.", styles['Normal']))
        
        elements.append(Spacer(1, 12))
    
    # Build the PDF
    doc.build(elements)
    
    return report_path

def generate_doc_report(results, report_id):
    """Generate DOC report with budget allocation results"""
    report_path = os.path.join(app.config['REPORT_FOLDER'], f"budget_report_{report_id}.docx")
    
    doc = docx.Document()
    
    # Add title
    doc.add_heading('UGC Budget Allocation Report', 0)
    
    # Add date
    doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph()
    
    # Add summary
    doc.add_heading(f"Total institutions analyzed: {len(results)}", level=2)
    
    # Create table for budget allocations
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    # Add header row
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Institution'
    header_cells[1].text = 'Score'
    header_cells[2].text = 'Budget Allocation (Rs.)'
    
    # Add data rows
    for result in results:
        row_cells = table.add_row().cells
        row_cells[0].text = result['name']
        row_cells[1].text = f"{result['total_score']}/10"
        row_cells[2].text = f"Rs. {result['budget_allocation']:,.2f}"
    
    doc.add_paragraph()
    
    # Add detailed recommendations for each institution
    doc.add_heading('Detailed Recommendations', level=2)
    
    for result in results:
        doc.add_heading(f"Institution: {result['name']}", level=3)
        
        if result['recommendations']:
            for recommendation in result['recommendations']:
                doc.add_paragraph(f"• {recommendation}", style='List Bullet')
        else:
            doc.add_paragraph("No specific recommendations.")
        
        doc.add_paragraph()
    
    # Save the document
    doc.save(report_path)
    
    return report_path

@app.route('/download/<filename>')
def download_report(filename):
    return send_file(os.path.join(app.config['REPORT_FOLDER'], filename), as_attachment=True)

if __name__ == '__main__':
    app.run(debug=True)